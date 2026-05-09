#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成格式规范的形势与政策课程作业 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── 页面设置 A4 ──
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin    = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin   = Cm(3.18)
sec.right_margin  = Cm(3.18)

def set_font(run, name_cn, name_en, size, bold=False):
    run.font.size = Pt(size)
    run.font.name = name_en
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_font(run, '黑体', 'SimHei', 22, bold=True)

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, '黑体', 'SimHei', 14, bold=True)

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, '宋体', 'SimSun', 12, bold=True)

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.85)  # 约2字符
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, '宋体', 'SimSun', 12)

def add_ref(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, '宋体', 'SimSun', 10.5)

# ════════ 正文 ════════
add_title('开放共赢 强贸兴邦')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run('——在高水平对外开放中迈向贸易强国')
set_font(run, '黑体', 'SimHei', 16)

# 引言
add_h1('引言')
add_body('当今世界正经历百年未有之大变局，国际经济格局深度调整，单边主义与保护主义沉渣泛起，全球产业链供应链面临重构压力。在这一复杂严峻的外部环境下，中国坚定不移推进高水平对外开放，加快建设贸易强国，不仅是实现中华民族伟大复兴的内在要求，更是推动构建人类命运共同体、促进世界经济复苏的大国担当。党的二十大报告明确提出"推进高水平对外开放"，党的二十届三中全会进一步强调"完善高水平对外开放体制机制"，为新时代我国对外开放事业指明了方向。深刻理解"开放共赢、强贸兴邦"的核心理念，对于我们准确把握国家战略大局、增强制度自信和道路自信具有重要意义。')

# 一
add_h1('一、高水平对外开放的时代背景与战略意义')
add_h2('（一）深刻变化的国际经贸格局')
add_body('进入二十一世纪第三个十年，全球贸易体系正在经历深刻重塑。一方面，世界贸易组织多边谈判进展迟缓，区域贸易协定快速发展，《区域全面经济伙伴关系协定》（RCEP）的生效实施标志着全球最大自贸区的诞生；另一方面，部分西方国家以"国家安全"为名推行"脱钩断链""小院高墙"策略，企图遏制中国的发展。据世界贸易组织统计，2025年全球货物贸易量增速放缓至2.7%，贸易摩擦频发使国际经贸环境的不确定性显著上升。面对这一形势，中国唯有以更高水平的开放应对外部挑战，才能在变局中开新局、于危机中育先机。')

add_h2('（二）从贸易大国到贸易强国的历史跨越')
add_body('改革开放四十多年来，中国已成为世界第一货物贸易大国。2025年，我国货物进出口总额达到43.85万亿元人民币，连续多年稳居全球第一。然而，"大"不等于"强"。贸易强国不仅要求规模领先，更要求在贸易结构、品牌价值、技术含量、规则制定能力和国际话语权等方面实现质的飞跃。当前，我国在高端制造、核心技术、国际品牌影响力等领域仍存在短板，服务贸易竞争力有待提升，参与全球经贸规则制定的深度和广度仍需拓展。因此，推进高水平对外开放、加快建设贸易强国，是我国经济从高速增长转向高质量发展的必然选择。')

add_h2('（三）高水平对外开放的战略定位')
add_body('高水平对外开放区别于一般意义上的开放，其核心特征体现在"制度型开放"上。它要求我们主动对接国际高标准经贸规则，稳步推进规则、规制、管理、标准等方面的制度性开放。这意味着不仅要降低关税壁垒、扩大市场准入，更要在知识产权保护、政府采购、数字贸易、环境标准等前沿领域与国际接轨。2024年，国务院发布的《关于加快内外贸一体化发展的若干措施》以及商务部出台的系列稳外贸政策，充分体现了制度型开放的战略导向。')

# 二
add_h1('二、迈向贸易强国的实践路径')
add_h2('（一）优化贸易结构，提升出口附加值')
add_body('建设贸易强国，首要任务是推动贸易结构从"以量取胜"向"以质取胜"转变。近年来，我国在新能源汽车、锂电池、光伏产品等领域实现了出口的跨越式增长，"新三样"出口额在2024年突破万亿元大关，成为我国外贸新的增长极。与此同时，跨境电商、海外仓等新业态新模式蓬勃发展，2025年一季度跨境电商进出口额同比增长超过11%，为中小企业参与国际竞争提供了新渠道。这些实践表明，依靠科技创新和产业升级驱动贸易发展，是迈向贸易强国的核心路径。')

add_h2('（二）深化自贸试验区和自由贸易港建设')
add_body('自贸试验区是中国推进高水平开放的试验田和排头兵。截至目前，全国已设立22个自贸试验区，形成了覆盖东西南北中的全方位开放格局。海南自由贸易港作为中国开放水平最高的区域，正在加快构建以零关税、低税率、简税制为特征的自由贸易政策体系。2025年，海南自贸港封关运作的各项准备工作有序推进，全岛封关后将成为全球最大的自由贸易港之一。这些制度创新实践为全国更大范围的开放积累了宝贵经验，也向世界展示了中国扩大开放的坚定决心。')

add_h2('（三）积极参与全球经济治理和规则制定')
add_body('贸易强国不仅是市场的参与者，更应成为规则的制定者。近年来，中国积极推动加入《全面与进步跨太平洋伙伴关系协定》（CPTPP）和《数字经济伙伴关系协定》（DEPA），主动对标国际最高标准开放规则。同时，中国在二十国集团、亚太经合组织、金砖国家合作机制等多边平台上，积极倡导开放型世界经济，反对贸易保护主义，为维护多边贸易体制贡献了中国智慧和中国方案。"一带一路"倡议提出十余年来，已与150多个国家和30多个国际组织签署合作文件，累计形成了超过3000个合作项目，有力促进了沿线国家的互联互通和共同发展。')

add_h2('（四）统筹发展与安全，构建开放安全新格局')
add_body('高水平对外开放绝非不顾风险的盲目开放，而是在开放中确保国家经济安全。我国在扩大开放的同时，不断完善外商投资安全审查制度，建立健全出口管制体系，加强关键领域的自主可控能力。特别是在粮食安全、能源安全、产业链安全等领域，坚持底线思维，确保开放进程可管可控。这种"放得开、管得住"的开放模式，体现了中国在推进高水平开放中的战略定力和治理智慧。')

# 三
add_h1('三、"开放共赢"理念的深刻内涵与世界意义')
add_body('"开放共赢"是中国推进高水平对外开放的核心理念，它超越了传统的零和博弈思维，彰显了中国作为负责任大国的胸怀与担当。中国国际进口博览会连续多届成功举办，累计意向成交额超过4200亿美元，生动诠释了"中国市场就是世界市场、中国机遇就是世界机遇"的开放理念。中国主动扩大进口，不仅满足了国内消费升级的需求，更为各国企业提供了广阔的市场空间，实现了互利共赢。')
add_body('从理论层面看，"开放共赢"理念根植于马克思主义政治经济学关于国际分工与世界市场的科学论述，同时汲取了中华优秀传统文化中"和合共生""天下大同"的思想精华，是马克思主义中国化时代化的重要理论成果。它回答了"在逆全球化浪潮下，世界经济向何处去"这一时代之问，为构建开放、包容、普惠、平衡、共赢的全球化新格局提供了中国方案。')

# 结语
add_h1('结语')
add_body('"强贸兴邦"是中华民族走向伟大复兴的必由之路，"开放共赢"是中国与世界共同繁荣的金钥匙。站在新的历史起点上，我们既要清醒认识到外部环境的严峻挑战，也要坚定信心于中国经济的强大韧性和巨大潜力。作为新时代的青年大学生，我们应当深入学习领会党中央关于高水平对外开放的战略部署，增强家国情怀和全球视野，努力成为具有国际竞争力的高素质人才，为推动中国从贸易大国迈向贸易强国贡献青春力量。唯有坚持开放合作、互利共赢，中国的发展之路才能越走越宽广，世界的未来也将因中国的开放而更加美好。')

# 参考文献
doc.add_paragraph()  # 空行
p = doc.add_paragraph()
run = p.add_run('参考文献：')
set_font(run, '黑体', 'SimHei', 12, bold=True)

refs = [
    '[1] 习近平.高举中国特色社会主义伟大旗帜 为全面建设社会主义现代化国家而团结奋斗——在中国共产党第二十次全国代表大会上的报告[R].北京：人民出版社，2022.',
    '[2] 中共中央关于进一步全面深化改革 推进中国式现代化的决定[M].北京：人民出版社，2024.',
    '[3] 国务院办公厅.关于加快内外贸一体化发展的若干措施[Z].2024.',
    '[4] 海关总署.2025年中国外贸进出口情况新闻发布会[EB/OL].2026.',
    '[5] 商务部.中国自由贸易试验区发展报告（2025）[R].2025.',
]
for r in refs:
    add_ref(r)

# 保存
out = '/Users/wangyijun/Desktop/形势与政策作业.docx'
doc.save(out)
print(f'✅ Word 文档已生成: {out}')
